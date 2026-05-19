"""
processor.py
============
Upgraded pipeline — zero Claude dependency, semantic similarity, real page fetching.

Changes from v1:
  - AI detection: Claude API → Hugging Face Inference API (free, purpose-built classifier)
  - Plagiarism: word overlap on snippets → semantic cosine similarity on full page content
  - Paraphrase + humanize: two DeepSeek calls → one merged call (same quality, half cost)
  - Internal similarity: word overlap → sentence embeddings via HF API
  - Table cells: same upgrades applied

Environment variables required:
    HF_API_KEY=...            ← free at huggingface.co/settings/tokens
    DEEPSEEK_API_KEY=...
    BRAVE_API_KEY=...
    INTERNAL_API_SECRET=...   ← used by main.py, not here
"""

import os
import re
import time
import requests
import numpy as np
from bs4 import BeautifulSoup
from openai import OpenAI
from docx import Document
from docx.shared import RGBColor

# ── Environment ───────────────────────────────────────────────────────────────

BRAVE_API_KEY    = os.getenv("BRAVE_API_KEY", "")
DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY", "")
HF_API_KEY       = os.getenv("HF_API_KEY", "")

HF_HEADERS = {"Authorization": f"Bearer {HF_API_KEY}"}

# Hugging Face model endpoints
HF_DETECTOR_URL  = "https://api-inference.huggingface.co/models/Hello-SimpleAI/chatgpt-detector-roberta"
HF_EMBEDDING_URL = "https://api-inference.huggingface.co/models/sentence-transformers/all-MiniLM-L6-v2"

# ── Thresholds ────────────────────────────────────────────────────────────────

AI_CONFIDENCE_MIN    = 0.75   # HF detector score to flag as AI-written
SEMANTIC_THRESHOLD   = 0.72   # cosine similarity to flag as plagiarism
INTERNAL_SIM_MIN     = 0.80   # internal doc similarity threshold (stricter)
MIN_WORDS            = 15     # skip very short paragraphs
PAGE_FETCH_TIMEOUT   = 8      # seconds to wait for a page fetch
MAX_PAGE_CHARS       = 6000   # chars of page content to embed (keeps cost low)


# ── Helpers ───────────────────────────────────────────────────────────────────

def _get_with_retry(url, headers, params=None, retries=2):
    for attempt in range(retries + 1):
        try:
            r = requests.get(url, headers=headers, params=params, timeout=10)
            r.raise_for_status()
            return r
        except Exception:
            if attempt == retries:
                return None
            time.sleep(1.5)


def _post_with_retry(url, headers, payload, retries=3):
    """POST to HF API with retry + model warm-up handling."""
    for attempt in range(retries + 1):
        try:
            r = requests.post(url, headers=headers, json=payload, timeout=30)
            if r.status_code == 503:
                # Model is loading — HF returns estimated_time
                wait = r.json().get("estimated_time", 20)
                time.sleep(min(wait, 30))
                continue
            r.raise_for_status()
            return r
        except Exception:
            if attempt == retries:
                return None
            time.sleep(2)


def is_reference_entry(text: str) -> bool:
    """Skip bibliography / reference list entries."""
    patterns = [
        r'^[A-Z][a-z]+,\s+[A-Z]',
        r'^\w.+\(\d{4}\)',
        r'^\w.+\d{4}\.',
        r'^https?://',
        r'^\[?\d+\]',
        r'^doi:',
        r'^\d+\.\s+[A-Z]',
        r'^[A-Z][a-z]+\s+[A-Z][a-z]+.*\d{4}',
        r'.*journal.*vol.*',
        r'.*pp\.\s*\d+',
    ]
    text_lower = text.lower()
    for pattern in patterns:
        if re.match(pattern, text.strip(), re.IGNORECASE):
            return True
    ref_keywords = ["doi:", "isbn", "issn", "vol.", "pp.", "et al.", "retrieved from"]
    return any(kw in text_lower for kw in ref_keywords)


# ── Step 1: Extract paragraphs ────────────────────────────────────────────────

def extract_paragraphs(filepath: str):
    """Return list of (doc_index, text) for paragraphs worth checking."""
    doc = Document(filepath)
    result = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text and len(text.split()) >= MIN_WORDS:
            result.append((i, text))
    return result


# ── Step 2: HF AI Detection ───────────────────────────────────────────────────

def _detect_ai_hf(text: str) -> float:
    """
    Call HF chatgpt-detector-roberta to score a paragraph.
    Returns float 0.0 (human) to 1.0 (AI).
    Falls back to 0.0 on failure — never blocks processing.
    """
    payload = {"inputs": text[:512]}  # model max context
    r = _post_with_retry(HF_DETECTOR_URL, HF_HEADERS, payload)
    if r is None:
        return 0.0
    try:
        data = r.json()
        # Response: [[{"label": "Fake", "score": 0.94}, {"label": "Real", "score": 0.06}]]
        if isinstance(data, list) and isinstance(data[0], list):
            for item in data[0]:
                if item.get("label", "").lower() in ("fake", "ai"):
                    return float(item["score"])
        return 0.0
    except Exception:
        return 0.0


def score_ai_likelihood(paragraphs: list) -> dict:
    """
    Score all paragraphs for AI likelihood via HF API.
    Returns {list_position: score}.
    """
    scores = {}
    for pos, (_, text) in enumerate(paragraphs):
        scores[pos] = _detect_ai_hf(text)
        time.sleep(0.2)  # light rate limiting
    return scores


# ── Step 3: Semantic Embeddings via HF ───────────────────────────────────────

def _get_embedding(text: str) -> list | None:
    """
    Get sentence embedding from HF all-MiniLM-L6-v2.
    Returns a list of floats, or None on failure.
    """
    payload = {"inputs": text[:512], "options": {"wait_for_model": True}}
    r = _post_with_retry(HF_EMBEDDING_URL, HF_HEADERS, payload)
    if r is None:
        return None
    try:
        data = r.json()
        # Returns a list of floats (the embedding vector)
        if isinstance(data, list) and isinstance(data[0], float):
            return data
        # Sometimes nested: [[float, ...]]
        if isinstance(data, list) and isinstance(data[0], list):
            return data[0]
        return None
    except Exception:
        return None


def _cosine_similarity(vec_a: list, vec_b: list) -> float:
    """Cosine similarity between two embedding vectors."""
    a = np.array(vec_a)
    b = np.array(vec_b)
    denom = np.linalg.norm(a) * np.linalg.norm(b)
    if denom == 0:
        return 0.0
    return float(np.dot(a, b) / denom)


def check_internal_similarity(paragraphs: list) -> set:
    """
    Flag paragraphs that are semantically too similar to each other.
    Uses sentence embeddings — catches paraphrased internal repetition.
    Falls back to word overlap if HF embedding fails.
    """
    flagged = set()
    texts = [t for _, t in paragraphs]

    # Try embedding approach first
    embeddings = []
    for text in texts:
        emb = _get_embedding(text)
        embeddings.append(emb)
        time.sleep(0.15)

    if all(e is not None for e in embeddings):
        for i in range(len(embeddings)):
            for j in range(i + 1, len(embeddings)):
                sim = _cosine_similarity(embeddings[i], embeddings[j])
                if sim >= INTERNAL_SIM_MIN:
                    flagged.add(i)
                    flagged.add(j)
    else:
        # Fallback: word overlap
        for i in range(len(texts)):
            for j in range(i + 1, len(texts)):
                wi = set(texts[i].lower().split())
                wj = set(texts[j].lower().split())
                if not wi or not wj:
                    continue
                overlap = len(wi & wj) / min(len(wi), len(wj))
                if overlap >= 0.60:
                    flagged.add(i)
                    flagged.add(j)

    return flagged


# ── Step 4: Plagiarism — Web + Academic ──────────────────────────────────────

def _fetch_page_text(url: str) -> str:
    """
    Fetch actual page content and strip to clean text.
    Returns up to MAX_PAGE_CHARS characters.
    Falls back to empty string on any error.
    """
    try:
        r = requests.get(
            url,
            timeout=PAGE_FETCH_TIMEOUT,
            headers={
                "User-Agent": (
                    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                    "AppleWebKit/537.36 (KHTML, like Gecko) "
                    "Chrome/124.0.0.0 Safari/537.36"
                )
            },
        )
        soup = BeautifulSoup(r.text, "html.parser")
        # Remove boilerplate tags
        for tag in soup(["script", "style", "nav", "footer", "header", "aside", "form"]):
            tag.decompose()
        text = soup.get_text(separator=" ", strip=True)
        # Collapse whitespace
        text = re.sub(r"\s+", " ", text)
        return text[:MAX_PAGE_CHARS]
    except Exception:
        return ""


def _semantic_match(para_text: str, source_text: str) -> float:
    """
    Compare paragraph against source text using semantic similarity.
    Splits source into chunks, scores each against paragraph, returns best score.
    Falls back to word overlap if embeddings fail.
    """
    if not source_text.strip():
        return 0.0

    para_emb = _get_embedding(para_text)
    if para_emb is None:
        # Fallback: word overlap
        para_words = set(para_text.lower().split())
        src_words  = set(source_text.lower().split())
        if not src_words:
            return 0.0
        return len(para_words & src_words) / len(para_words)

    # Split source into overlapping chunks of ~300 chars for better coverage
    chunk_size = 300
    step       = 150
    chunks     = [
        source_text[i : i + chunk_size]
        for i in range(0, len(source_text), step)
        if len(source_text[i : i + chunk_size].split()) >= 10
    ]

    best_score = 0.0
    for chunk in chunks[:20]:  # cap at 20 chunks to control HF calls
        chunk_emb = _get_embedding(chunk)
        if chunk_emb is None:
            continue
        score = _cosine_similarity(para_emb, chunk_emb)
        if score > best_score:
            best_score = score
        time.sleep(0.1)

    return best_score


def check_brave(text: str) -> tuple:
    """
    Search Brave for the paragraph, fetch actual page content,
    compare semantically. Returns (is_flagged, best_url, best_score).
    """
    query = text[:120].strip()
    headers = {
        "Accept": "application/json",
        "Accept-Encoding": "gzip",
        "X-Subscription-Token": BRAVE_API_KEY,
    }
    params = {"q": f'"{query}"', "count": 5}

    try:
        r = _get_with_retry(
            "https://api.search.brave.com/res/v1/web/search",
            headers=headers,
            params=params,
        )
        if r is None:
            return False, None, 0.0

        results = r.json().get("web", {}).get("results", [])
        if not results:
            return False, None, 0.0

        best_score, best_url = 0.0, None

        for result in results[:3]:  # check top 3 URLs
            url         = result.get("url", "")
            page_text   = _fetch_page_text(url)

            if not page_text:
                # Fallback to snippet if page fetch fails
                page_text = result.get("description", "") + " " + result.get("title", "")

            score = _semantic_match(text, page_text)
            if score > best_score:
                best_score = score
                best_url   = url

        return best_score >= SEMANTIC_THRESHOLD, best_url, round(best_score, 2)

    except Exception:
        return False, None, 0.0


def check_openalex(text: str) -> tuple:
    """
    Check paragraph against OpenAlex academic database.
    Uses semantic similarity against title + abstract.
    """
    query = text[:100].strip()
    try:
        r = _get_with_retry(
            "https://api.openalex.org/works",
            headers={},
            params={"search": query, "per-page": 5},
        )
        if r is None:
            return False, None, 0.0

        results = r.json().get("results", [])
        if not results:
            return False, None, 0.0

        best_score, best_url = 0.0, None

        for work in results:
            title        = work.get("title", "") or ""
            abstract_inv = work.get("abstract_inverted_index", {}) or {}
            abstract     = " ".join(abstract_inv.keys())
            source_text  = (title + " " + abstract).strip()

            if not source_text:
                continue

            score = _semantic_match(text, source_text)
            if score > best_score:
                best_score = score
                best_url   = work.get("id", "OpenAlex record")

        return best_score >= SEMANTIC_THRESHOLD, best_url, round(best_score, 2)

    except Exception:
        return False, None, 0.0


# ── Step 5: DeepSeek — Merged Paraphrase + Humanize ─────────────────────────

def _get_deepseek_client():
    return OpenAI(api_key=DEEPSEEK_API_KEY, base_url="https://api.deepseek.com")


def paraphrase_and_humanize(text: str, client, hard_mode: bool = False) -> str:
    """
    Single DeepSeek call that paraphrases AND humanizes in one pass.
    Replaces the old paraphrase() + humanize() two-call pattern.
    hard_mode=True for paragraphs with very high web similarity (≥0.85).
    """
    if hard_mode:
        instruction = (
            "This paragraph is almost identical to a web source. "
            "Rewrite it COMPLETELY from scratch — different structure, different vocabulary, "
            "different logical flow. Imagine you are explaining this idea to a colleague "
            "entirely in your own words. Nothing from the original phrasing should survive."
        )
    else:
        instruction = (
            "Rewrite the following paragraph so it is completely original "
            "and passes plagiarism and AI detection tools."
        )

    prompt = (
        f"{instruction}\n\n"
        "Apply ALL of the following in a single pass:\n\n"
        "ORIGINALITY:\n"
        "1. Completely restructure every sentence — flip subject/object, split or merge clauses\n"
        "2. Replace key terms with precise, high-quality synonyms\n"
        "3. Reorder supporting ideas where the logic still holds\n"
        "4. Preserve ALL technical terms, numbers, statistics, and facts exactly\n\n"
        "HUMAN VOICE:\n"
        "5. Mix sentence lengths — some short and punchy, some long and complex\n"
        "6. Alternate between active and passive voice naturally\n"
        "7. Add 1-2 natural transitional phrases a real writer would use\n"
        "8. Introduce minor stylistic asymmetry — not every sentence the same pattern\n"
        "9. Slightly loosen 1-2 overly formal phrasings without losing academic register\n\n"
        "OUTPUT:\n"
        "- Return ONLY the rewritten paragraph\n"
        "- No preamble, no commentary, no quotation marks around the output\n"
        "- Same approximate length as the original\n\n"
        f"Original paragraph:\n{text}"
    )

    try:
        response = client.chat.completions.create(
            model="deepseek-chat",
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are a senior academic editor who rewrites content to be "
                        "completely original and naturally human-sounding. "
                        "You output only the rewritten paragraph with zero extra text."
                    ),
                },
                {"role": "user", "content": prompt},
            ],
            max_tokens=1024,
            temperature=0.88,
        )
        return response.choices[0].message.content.strip()
    except Exception:
        return text  # safe fallback — return original if DeepSeek fails


# ── Step 6: Replace paragraph text in .docx ──────────────────────────────────

def replace_paragraph_text(para, new_text: str):
    """Replace paragraph text preserving original formatting, highlight green."""
    if para.runs:
        first = para.runs[0]
        bold, italic, underline = first.bold, first.italic, first.underline
        font_name, font_size    = first.font.name, first.font.size
    else:
        bold = italic = underline = False
        font_name = font_size = None

    for run in para.runs:
        run._element.getparent().remove(run._element)

    new_run           = para.add_run(new_text)
    new_run.bold      = bold
    new_run.italic    = italic
    new_run.underline = underline
    if font_name:
        new_run.font.name = font_name
    if font_size:
        new_run.font.size = font_size
    new_run.font.color.rgb = RGBColor(0, 130, 0)  # green = edited


# ── Main pipeline ─────────────────────────────────────────────────────────────

def process_document(input_path: str, output_path: str) -> dict:
    """
    Full pipeline. Returns a report dict for the API response.
    """
    paragraphs = extract_paragraphs(input_path)
    if not paragraphs:
        raise ValueError("No processable paragraphs found in document.")

    total_paragraphs = len(paragraphs)

    # ── HF AI detection ──────────────────────────────────────────────────────
    ai_scores = score_ai_likelihood(paragraphs)
    ai_flagged_pos = {
        pos for pos, score in ai_scores.items()
        if score >= AI_CONFIDENCE_MIN
    }

    # ── Internal semantic similarity ──────────────────────────────────────────
    internally_flagged = check_internal_similarity(paragraphs)

    # ── Web + academic plagiarism ─────────────────────────────────────────────
    web_flagged      = {}
    academic_flagged = {}

    for idx, (para_idx, text) in enumerate(paragraphs):
        if is_reference_entry(text):
            continue

        is_web, url, score = check_brave(text)
        if is_web:
            web_flagged[para_idx] = (url, score)
        time.sleep(0.3)

        is_acad, acad_url, acad_score = check_openalex(text)
        if is_acad:
            academic_flagged[para_idx] = (acad_url, acad_score)
        time.sleep(0.2)

    # ── Build master flag set ─────────────────────────────────────────────────
    all_flagged  = set()
    skipped_refs = []

    for list_pos, (para_idx, text) in enumerate(paragraphs):
        reasons = []
        if list_pos in ai_flagged_pos:
            reasons.append("ai")
        if list_pos in internally_flagged:
            reasons.append("internal")
        if para_idx in web_flagged:
            reasons.append("web")
        if para_idx in academic_flagged:
            reasons.append("academic")

        if reasons:
            if is_reference_entry(text):
                skipped_refs.append(para_idx)
            else:
                all_flagged.add(para_idx)

    # ── Paraphrase + humanize (merged, single DeepSeek call per paragraph) ────
    client            = _get_deepseek_client()
    doc               = Document(input_path)
    paraphrased_count = 0
    report_items      = []

    for i, para in enumerate(doc.paragraphs):
        if i not in all_flagged or not para.text.strip():
            continue

        original  = para.text.strip()

        # Use hard mode if web similarity was very high
        hard = i in web_flagged and web_flagged[i][1] >= 0.85

        rewritten = paraphrase_and_humanize(original, client, hard_mode=hard)
        time.sleep(0.8)

        replace_paragraph_text(para, rewritten)
        paraphrased_count += 1

        flags    = []
        list_pos = next(
            (lp for lp, (pi, _) in enumerate(paragraphs) if pi == i), None
        )
        if list_pos is not None and list_pos in ai_flagged_pos:
            confidence = int(ai_scores.get(list_pos, 0) * 100)
            flags.append(f"AI-written (detector confidence: {confidence}%)")
        if list_pos is not None and list_pos in internally_flagged:
            flags.append("Internal repetition (semantic)")
        if i in web_flagged:
            url, score = web_flagged[i]
            flags.append(f"Web match {int(score * 100)}% — {url}")
        if i in academic_flagged:
            url, score = academic_flagged[i]
            flags.append(f"Academic match {int(score * 100)}% — {url}")

        report_items.append({
            "paragraph": i + 1,
            "preview":   original[:100] + ("..." if len(original) > 100 else ""),
            "flags":     flags,
            "action":    "paraphrased",
        })

    for para_idx in skipped_refs:
        report_items.append({
            "paragraph": para_idx + 1,
            "preview":   "",
            "flags":     ["Matched source — skipped (reference entry)"],
            "action":    "skipped",
        })

    # ── Process table cells ───────────────────────────────────────────────────
    tables_processed = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if not text or len(text.split()) < MIN_WORDS:
                        continue
                    if is_reference_entry(text):
                        continue

                    is_web, _, web_score   = check_brave(text)
                    is_acad, _, acad_score = check_openalex(text)
                    time.sleep(0.2)

                    if is_web or is_acad:
                        hard      = web_score >= 0.85
                        rewritten = paraphrase_and_humanize(text, client, hard_mode=hard)
                        time.sleep(0.8)
                        replace_paragraph_text(para, rewritten)
                        paraphrased_count += 1
                        tables_processed  += 1

    doc.save(output_path)

    return {
        "total_paragraphs_checked": total_paragraphs,
        "paragraphs_paraphrased":   paraphrased_count,
        "tables_cells_rewritten":   tables_processed,
        "references_skipped":       len(skipped_refs),
        "items":                    report_items,
    }


# ── Quick analysis (no AI calls) ──────────────────────────────────────────────

def analyse_document(filepath: str) -> dict:
    """
    Fast scan of a .docx file.
    Returns word count, paragraph count, table count, image count.
    No AI calls — runs in under a second.
    """
    doc = Document(filepath)

    word_count      = 0
    paragraph_count = 0
    table_count     = len(doc.tables)
    image_count     = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            word_count      += len(text.split())
            paragraph_count += 1

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if text:
                        word_count += len(text.split())

    try:
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                image_count += 1
    except Exception:
        pass

    return {
        "word_count":      word_count,
        "paragraph_count": paragraph_count,
        "table_count":     table_count,
        "image_count":     image_count,
    }
